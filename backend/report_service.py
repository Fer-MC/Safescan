"""
report_service.py — Generación del informe de inspección en Word (python-docx).

Produce un INFORME DE INSPECCIÓN VISUAL DE APOYO, multiidioma (es/en/ca).
La EMPRESA USUARIA es la protagonista del documento; VETLLA aparece en el header
de todas las páginas como "Generado con VETLLA".

Características de formato:
  - Tipografía Aptos (cuerpo) / Aptos Display (títulos), declarada en el documento.
  - Foto analizada incrustada, centrada y COMPRIMIDA localmente (sin llamadas API).
  - Tablas: centrado vertical de celdas, spacing 6pt; solo la celda de
    "Clasificación" lleva sombreado pastel (ShadingType CLEAR).
  - Saltos de página estratégicos para evitar cortes de contenido.
"""

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from i18n import CLASIF_LABEL, CONF_LABEL, DEADLINE_LABEL, normalize_lang, report_strings

try:
    from PIL import Image, ImageOps
    _PIL_OK = True
except ImportError:
    _PIL_OK = False

# --- Paleta VETLLA ---
AZUL    = RGBColor(0x1B, 0x2F, 0x4E)
NARANJA = RGBColor(0xE8, 0x61, 0x1A)
GRIS    = RGBColor(0x4A, 0x55, 0x68)

# Fuentes
FONT_BODY = "Aptos"
FONT_TITLE = "Aptos Display"

# Sombreado SOLO para la celda de clasificación.
# "critical" usa un rojo sólido con texto blanco para que destaque claramente
# frente a "important" (que se queda en el tono pastel). El resto, sin cambios.
COLOR_CLASIF = {
    "critical":  "D85A30",
    "important": "F5D6D6",
    "minor":     "FBE0C7",
    "conform":   "DCEBD1",
}
# Color de texto de la celda de clasificación, por clave (por defecto GRIS oscuro/negro
# del documento; "critical" es la única que necesita texto blanco sobre su fondo sólido).
TEXT_CLASIF = {
    "critical": RGBColor(0xFF, 0xFF, 0xFF),
}

# Ancho de columna etiqueta
COL1_CM = 4.5
PAGE_CONTENT_CM = 16.0

# Compresión de imagen
IMG_MAX_PX = 1400
IMG_QUALITY = 72
IMG_DOC_WIDTH_CM = 12.0


def _set_document_fonts(doc):
    """Fija Aptos como fuente por defecto del documento."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_BODY)
    rfonts.set(qn("w:hAnsi"), FONT_BODY)
    rfonts.set(qn("w:cs"), FONT_BODY)


def _add_header(doc, text):
    """Añade texto al header del documento (todas las páginas)."""
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run(text)
        header_run.font.size = Pt(8)
        header_run.font.color.rgb = GRIS
        header_run.font.name = FONT_BODY


def _title(doc, text, size=16, color=AZUL, align=WD_ALIGN_PARAGRAPH.LEFT,
           space_before=10, font=FONT_TITLE, bold=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return p


def _page_break(doc):
    """Añade un salto de página."""
    doc.add_page_break()

def _prevent_row_split(row):
    """Impide que una fila de tabla se divida entre dos páginas."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _keep_with_next(paragraph):
    """Mantiene el párrafo pegado al siguiente (no se separan por salto de página)."""
    p_fmt = paragraph.paragraph_format
    p_fmt.keep_with_next = True



def _compress_image(image_bytes: bytes) -> BytesIO | None:
    """Comprime la imagen localmente (sin API)."""
    if not _PIL_OK or not image_bytes:
        return None
    try:
        im = Image.open(BytesIO(image_bytes))
        im = ImageOps.exif_transpose(im)
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        im.thumbnail((IMG_MAX_PX, IMG_MAX_PX), Image.LANCZOS)
        out = BytesIO()
        im.save(out, format="JPEG", quality=IMG_QUALITY, optimize=True)
        out.seek(0)
        return out
    except Exception:
        return None


def generar_informe(
    analisis: dict,
    *,
    empresa: str = "",
    centro: str = "",
    responsable: str = "",
    resp_zona: str = "",
    supervisor: str = "",
    lang: str = "es",
    image_bytes: bytes | None = None,
) -> BytesIO:
    lang = normalize_lang(lang)
    T = report_strings(lang)
    clabel = CLASIF_LABEL[lang]
    conflabel = CONF_LABEL[lang]
    deadlinelabel = DEADLINE_LABEL[lang]

    doc = Document()
    _set_document_fonts(doc)

    # --- Header en todas las páginas ---
    _add_header(doc, T["brand_seal"])

    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    col2_cm = PAGE_CONTENT_CM - COL1_CM

    # --- Título del informe (empresa como protagonista) ---
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = titulo.add_run(empresa.strip() or T["app_subtitle"])
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = AZUL
    tr.font.name = FONT_TITLE

    # --- Subtítulo en MAYÚSCULAS ---
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(T["app_subtitle"].upper())
    sr.font.size = Pt(11)
    sr.font.color.rgb = GRIS
    sr.font.name = FONT_BODY

    # --- Bloque de datos iniciales (sin Empresa, con barras |) ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lineas = [f"{T['date']}: {datetime.now().strftime('%d/%m/%Y %H:%M')}"]
    if centro:
        lineas.append(f"{T['site']}: {centro}")
    if responsable:
        lineas.append(f"{T['inspector']}: {responsable}")
    mr = meta.add_run(" | ".join(lineas))
    mr.font.size = Pt(9)
    mr.font.name = FONT_BODY

    # --- Foto analizada (centrada, comprimida) ---
    compressed = _compress_image(image_bytes)
    if compressed is not None:
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.paragraph_format.space_before = Pt(10)
        run_pic = pic_p.add_run()
        run_pic.add_picture(compressed, width=Cm(IMG_DOC_WIDTH_CM))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(T["image_caption"])
        cr.font.size = Pt(8)
        cr.bold = True
        cr.font.color.rgb = GRIS
        cr.font.name = FONT_BODY

    # --- Resumen (se queda en la primera hoja, junto a la imagen si cabe) ---
    _title(doc, T["summary_heading"], size=14, space_before=10)
    doc.add_paragraph(analisis.get("resumen") or T["no_summary"])
    n = analisis.get("num_observaciones", len(analisis.get("observaciones", [])))
    cnt = doc.add_paragraph()
    cnt_run = cnt.add_run(f"{T['obs_count']}: {n}")
    cnt_run.bold = True

    # --- Observaciones (empiezan en hoja nueva) ---
    _page_break(doc)
    _title(doc, T["obs_heading"], size=14, space_before=12)
    observaciones = analisis.get("observaciones", [])
    if not observaciones:
        doc.add_paragraph(T["no_obs"])
    else:
        for i, obs in enumerate(observaciones, 1):
            titulo_obs = _title(doc, f"{i}. {obs.get('categoria', '')}", size=12,
                   space_before=10, font=FONT_TITLE)
            _keep_with_next(titulo_obs)

            tabla = doc.add_table(rows=0, cols=2)
            tabla.style = "Table Grid"

            clave = obs.get("clasificacion", "important")

            def fila(k, v, shade=None, text_color=None, text_bold=False):
                row = tabla.add_row()
                _prevent_row_split(row)
                cells = row.cells
                cells[0].width = Cm(COL1_CM)
                cells[1].width = Cm(col2_cm)
                for c in cells:
                    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p0 = cells[0].paragraphs[0]
                p0.paragraph_format.space_before = Pt(6)
                r0 = p0.add_run(k)
                r0.bold = True
                r0.font.size = Pt(9)
                r0.font.name = FONT_BODY
                p1 = cells[1].paragraphs[0]
                p1.paragraph_format.space_before = Pt(6)
                r1 = p1.add_run(str(v))
                r1.font.size = Pt(9)
                r1.font.name = FONT_BODY
                if text_color:
                    r1.font.color.rgb = text_color
                if text_bold:
                    r1.bold = True
                if shade:
                    _shade_cell(cells[1], shade)
                return cells

            fila(T["f_desc"], obs.get("descripcion", ""))
            if obs.get("ubicacion"):
                fila(T["f_loc"], obs.get("ubicacion", ""))
            fila(T["f_class"], clabel.get(clave, clave), shade=COLOR_CLASIF.get(clave),
                 text_color=TEXT_CLASIF.get(clave), text_bold=(clave == "critical"))
            if obs.get("acciones"):
                fila(T["f_actions"], "\n".join(f"• {a}" for a in obs.get("acciones", [])))
            # Plazo estimado: NO viene de la IA, se deriva por regla fija de la clasificación.
            fila(T["f_deadline"], deadlinelabel.get(clave, deadlinelabel["important"]))
            fila(T["f_norm"], ", ".join(obs.get("normativa", [])) or "—")
            fila(T["f_conf"], conflabel.get(obs.get("confianza", "medium")))

    # --- Limitaciones (pequeñas y en cursiva) ---
    _title(doc, T["limits_heading"], size=12, space_before=12)
    for lim in analisis.get("limitaciones", []) or [T["no_limits"]]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(lim)
        run.font.size = Pt(8)
        run.italic = True
        run.font.name = FONT_BODY

    # --- Aviso legal + firma ---
    p_legal = doc.add_paragraph()
    p_legal.paragraph_format.space_before = Pt(20)
    rl = p_legal.add_run(T["legal"])
    rl.font.size = Pt(7.5)
    rl.italic = True
    rl.font.color.rgb = GRIS
    rl.font.name = FONT_BODY

    firma = doc.add_paragraph()
    firma.paragraph_format.space_before = Pt(14)
    fr = firma.add_run(T["sign_line"])
    fr.font.name = FONT_BODY

    # Firmantes adicionales opcionales — solo aparecen si el campo se rellenó en el formulario.
    if resp_zona.strip():
        firma_zona = doc.add_paragraph()
        firma_zona.paragraph_format.space_before = Pt(8)
        fz = firma_zona.add_run(T["sign_zone"])
        fz.font.name = FONT_BODY

    if supervisor.strip():
        firma_sup = doc.add_paragraph()
        firma_sup.paragraph_format.space_before = Pt(8)
        fs = firma_sup.add_run(T["sign_supervisor"])
        fs.font.name = FONT_BODY

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _shade_cell(cell, hex_color: str):
    """Sombreado de celda con ShadingType CLEAR."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)
